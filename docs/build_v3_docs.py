from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1D2939"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = "087A55"
AMBER = "B54708"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
TODAY = date.today().isoformat()


def set_font(run, size=11, bold=None, color=INK, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def style_callout_paragraph(paragraph, fill=LIGHT_GRAY, border="D0D5DD", center=False):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "start", "bottom", "end"):
        node = p_bdr.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            p_bdr.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), border)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_font(run, 9, color=MUTED)


def configure_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run(running_title), 9, bold=True, color=MUTED)
    set_font(hp.add_run("  |  v3.0"), 9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    set_font(fp.add_run("紫东太初 · Token API 成本测算  |  "), 9, color=MUTED)
    add_page_number(fp)


def add_cover(doc, kicker, title, subtitle):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run(kicker.upper()), 11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run(title), 28, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_font(p.add_run(subtitle), 14, color=MUTED)
    cp = doc.add_paragraph()
    style_callout_paragraph(cp, LIGHT_GRAY, center=True)
    set_font(cp.add_run("理论上限 → 工程预设 → 严格匹配压测 → 成本与报价"), 11, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("版本 v3.0"), 11, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f"更新日期：{TODAY}"), 10, color=MUTED)
    doc.add_page_break()


def add_intro_callout(doc, title, text, fill=LIGHT_GRAY, color=DARK_BLUE):
    p = doc.add_paragraph()
    style_callout_paragraph(p, fill)
    set_font(p.add_run(title + "："), 11, bold=True, color=color)
    set_font(p.add_run(text), 11, color=INK)


def add_step(doc, number, title, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(10)
    set_font(p.add_run(f"{number}. {title}"), 13, bold=True, color=BLUE)
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(7)


def add_formula(doc, label, formula, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label), 11, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(formula), 10.5, bold=True, color=INK, name="Consolas")
    if note:
        p = doc.add_paragraph(note)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(7)
        set_font(p.runs[0], 10, color=MUTED)


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), 10, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), 10, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build_usage_guide():
    doc = Document()
    configure_document(doc, "Token API 成本测算工具 · 使用指南")
    add_cover(doc, "Operator Guide", "Token API 成本测算工具", "使用指南 · v3.1 硬件池、完整副本与压测校准")

    doc.add_heading("1. 工具定位", level=1)
    add_intro_callout(doc, "核心口径", "理论上限只用于容量边界；无压测时使用工程场景预设和 Amdahl 非线性扩展；存在严格匹配压测时，单机整机 TPS 或多机完整副本 TPS 自动覆盖估算。")
    doc.add_paragraph("本工具用于估算模型部署所需资源、Aggregate Output TPS、月输出 Token、每百万输出 Token 的硬件成本及建议报价。结果用于容量规划，不替代正式 SLA 压测。")

    doc.add_heading("2. 快速使用流程", level=1)
    steps = [
        ("选择模型", "从模型分类中选择预设，或手动填写模型名称、总参数和每 Token 激活参数，并保存为浏览器本地预设。"),
        ("选择权重精度", "BF16/FP16、FP8、INT4/MXFP4 影响权重显存；量化算力倍率没有实测时保持 1.00。"),
        ("选择硬件", "优先选择已核实的整机预设。Ascend 910C 预设锁定为 8 卡/节点；选择自定义整机后，必须重新确认规格和合同月租。"),
        ("选择部署模式", "单机填写当前机器实际卡数；多机分别填写单副本拓扑节点数、单副本实际总卡数和规划副本数。"),
        ("选择硬件扩容方式", "默认选择“同步增加硬件”；也可使用固定节点池。只有部署或压测已验证整张空卡可共用时，才在高级设置中填写每组可运行副本数。"),
        ("选择估算模板", "无压测时优先使用“在线均衡”；交互低并发和吞吐优先用于敏感性比较。"),
        ("填写计费", "根据合同选择按整节点或按实际卡数计费，并填写对应月租。"),
        ("选择测算模式", "容量模式计算现有配置的可持续产能；QPS 模式根据业务需求自动反算服务副本。"),
        ("检查有效性", "先确认权重显存、单机卡数上限、单副本卡分布、付费池容量、拓扑和计费配置均通过校验。"),
        ("读取成本", "查看总月租、付费节点、物理卡容量、副本分配卡数、空余卡数、月输出 Token 与建议报价。"),
        ("用压测校准", "单机保存当前整机 TPS，多机保存一个完整副本 TPS；补齐对应拓扑、运行时和工作负载指纹后应用。"),
    ]
    for idx, (title, text) in enumerate(steps, 1):
        add_step(doc, idx, title, text)

    doc.add_heading("3. 三类估算模板", level=1)
    add_data_table(
        doc,
        ["模板", "内核与服务效率", "节点内不可并行", "跨节点不可并行", "用途"],
        [
            ["交互低并发", "5%", "15%", "25%", "低并发、时延优先"],
            ["在线均衡", "10%", "10%", "15%", "默认规划起点"],
            ["吞吐优先", "20%", "5%", "10%", "高并发、批处理充分"],
        ],
        [1800, 1800, 1800, 1800, 2160],
    )
    doc.add_paragraph("这些百分比是工程规划先验，不是厂商性能承诺。修改任一效率参数后，页面自动切换为“自定义参数”。")

    doc.add_heading("4. 单机与多机的区别", level=1)
    add_data_table(
        doc,
        ["概念", "含义", "对吞吐的作用", "对租金的作用"],
        [
            ["单节点标准卡数", "所选硬件 SKU 的物理规格", "不直接作为吞吐倍率", "决定节点容量与整机价格口径"],
            ["单副本实际总卡数", "一个完整模型副本实际占用的卡", "按节点分布做非线性估算", "决定物理整卡容量上限"],
            ["单副本拓扑节点数", "共同承载一个完整模型副本的节点", "引入跨节点通信效率", "不是付费池总节点数"],
            ["付费硬件池", "已经购买或整租的节点与卡容量", "不直接增加单副本吞吐", "直接决定总月租"],
            ["完整副本数", "可独立处理请求的完整模型实例", "按副本扩展并扣多副本调度损耗", "决定需要增加多少硬件或资源组"],
        ],
        [1700, 2860, 2600, 2200],
    )
    add_intro_callout(doc, "判定规则", "单机实际卡数不得超过硬件 SKU 的单节点标准卡数；超限后停止成本和报价并提示切换多机。多机先定义单副本占用，再把副本放入付费硬件池。")

    doc.add_heading("5. 容量模式", level=1)
    doc.add_paragraph("容量模式回答“当前付费资源在给定规划口径下，每月能够输出多少 Token”。")
    add_formula(doc, "总吞吐", "单副本吞吐 × 完整副本数 × 多副本调度效率")
    add_formula(doc, "月输出 Token", "总吞吐 × 运行天数 × 日均有效负载小时 × 3600")
    doc.add_paragraph("结果区同时显示交互低并发、当前模板和吞吐优先三个场景，便于观察参数敏感性。")

    doc.add_heading("6. QPS 模式", level=1)
    doc.add_paragraph("QPS 模式回答“为了满足指定业务请求量，需要付费多少个完整副本”。")
    add_formula(doc, "输出需求", "QPS × 平均输出 Token")
    add_formula(doc, "无压测等效规划负载", "QPS ×（输出 Token + 输入 Token × Prefill 等效负载系数）")
    add_formula(doc, "服务副本数", "向上取整[等效规划负载 ÷（单副本规划能力 × 目标负载率 × 调度效率）]")
    add_formula(doc, "应付费副本", "服务副本 + 冗余副本")
    add_intro_callout(doc, "计费边界", "Prefill 等效负载只用于反算资源；单位成本的产量分母始终使用实际输出 Token。冗余副本产生租金，但不计入常态业务输出。", fill="FFF4E5", color=AMBER)

    doc.add_heading("7. 硬件扩容与计费方式", level=1)
    add_data_table(
        doc,
        ["副本增加时硬件如何变化", "付费节点计算", "成本变化"],
        [
            ["同步增加硬件（默认）", "副本数 × 单副本节点数", "月租随副本数同步增长"],
            ["使用固定节点池", "先算每组可运行副本数，再算需要几组", "月租按资源组阶梯增长"],
            ["高级：实测空卡共用", "采用实测的每组可运行副本数", "实测容量内月租不变，跨容量增加资源组"],
        ],
        [2200, 3960, 3200],
    )
    add_formula(doc, "同步增加硬件", "总节点数 = 副本数 × 单副本节点数")
    add_formula(doc, "固定节点池每组可运行副本数", "向下取整[每组付费节点数 ÷ 单副本节点数]")
    add_formula(doc, "固定节点池资源组数", "向上取整[规划副本数 ÷ 每组可运行副本数]")
    add_formula(doc, "高级实测空卡共用", "手工回填同配置部署或压测确认值，且不得超过向下取整[每组物理卡容量 ÷ 单副本实际总卡数]")
    add_intro_callout(doc, "空卡边界", "空卡仅指未分配给任何副本的整张卡；已分配卡上的剩余显存不属于空卡。只有 NUMA、互联拓扑、进程隔离和调度均已通过部署或压测验证，才能填写每组实测可运行副本数。", fill="FFF4E5", color=AMBER)
    add_data_table(
        doc,
        ["方式", "总月租计算", "典型现象"],
        [
            ["按整节点", "节点月租 × 付费节点数", "同一已付费节点内增加启用卡数，租金可不变"],
            ["按实际卡数", "单卡月租 × 计费卡数", "固定节点池按预留卡数；同步增加硬件按副本实际卡数"],
        ],
        [1900, 3100, 4360],
    )
    add_formula(doc, "硬件成本", "总月租 ÷ 月输出 Token × 1,000,000")
    add_formula(doc, "建议报价", "每百万 Token 硬件成本 ÷（1 - 目标毛利率）")
    doc.add_paragraph("建议报价只覆盖硬件租金，正式价格还应叠加电力、网络、存储、运维、许可和税费。")

    doc.add_heading("8. 压测录入与复用", level=1)
    doc.add_paragraph("压测数据保存于当前浏览器，不上传服务器。单机保存当前整机 Aggregate Output TPS；多机保存一个完整模型副本的 Aggregate Output TPS。只有以下关键指纹匹配时，系统才会应用实测 TPS：")
    add_data_table(
        doc,
        ["类别", "必须匹配的内容"],
        [
            ["模型", "模型、结构、总参数、激活参数、权重精度"],
            ["硬件", "卡型、基准算力、量化倍率、单节点标准卡数、单副本实际总卡数、单副本节点数"],
            ["拓扑", "TP、PP、EP；多机额外匹配互联类型、互联带宽、RDMA"],
            ["运行时", "推理框架、框架或镜像版本"],
            ["工作负载", "ISL、OSL、并发，以及已设置的 P95 TTFT/ITL"],
        ],
        [1900, 7460],
    )
    doc.add_paragraph("任一关键字段变化，旧压测自动失配并回退到工程估算。")

    doc.add_heading("9. 使用边界", level=1)
    add_intro_callout(doc, "权重显存", "只校验权重是否能放入扣除预留后的显存，不代表 KV Cache、激活、通信缓冲和运行时一定满足。")
    add_intro_callout(doc, "理论吞吐", "FLOPS 公式是 Decode 计算容量上限，不是生产 Aggregate Output TPS。")
    add_intro_callout(doc, "正式交付", "上线容量、SLA 和正式报价必须使用同模型、同拓扑、同运行时和同工作负载的压测数据。")

    path = OUT / "Token_API_成本测算工具_使用指南_v3.1.docx"
    doc.save(path)
    return path


def build_logic_report():
    doc = Document()
    configure_document(doc, "Token API 成本测算工具 · 计算口径与更新报告")
    add_cover(doc, "Calculation Specification", "Token API 成本测算工具", "计算口径与更新报告 · v3.1")

    doc.add_heading("1. 本次更新结论", level=1)
    add_intro_callout(doc, "主要变化", "多机成本入口合并为一个问题：“副本增加时，硬件如何变化？”默认按副本同步增加硬件；固定节点池按节点容量分组；实测空卡共用收进高级设置。无压测吞吐继续使用 Amdahl 非线性估算，严格匹配压测仍拥有最高优先级。")
    doc.add_paragraph("此次更新解决“节点、实际卡数、完整副本和付费资源混为同一数量”的问题。系统现在分别计算单副本吞吐、硬件池容量、付费资源组、总月租和单位 Token 成本。")

    doc.add_heading("2. 计算优先级", level=1)
    add_data_table(
        doc,
        ["优先级", "数据口径", "系统行为"],
        [
            ["1", "严格匹配实测", "单机采用整机 TPS；多机采用单副本 TPS"],
            ["2", "工程场景预设", "采用 Amdahl 多卡/多机扩展和效率参数"],
            ["3", "理论上限", "仅用于容量边界、利用率和合理性检查"],
        ],
        [1200, 3000, 5160],
    )

    doc.add_heading("3. 模型与显存", level=1)
    add_formula(doc, "权重显存", "总参数量（B）× 每参数字节数")
    add_data_table(
        doc,
        ["精度", "显存近似", "计算倍率口径"],
        [
            ["BF16/FP16", "2 B/参数", "基准倍率固定 1.00"],
            ["FP8", "1 B/参数", "默认 1.00，实测后保存"],
            ["INT4/MXFP4", "0.5 B/参数", "默认 1.00，实测后保存"],
        ],
        [2500, 2500, 4360],
    )
    add_formula(doc, "单副本可用显存", "单卡显存 × 单副本实际总卡数 ×（1 - 显存预留率）")
    doc.add_paragraph("总参数用于权重显存；MoE 的每 Token 激活参数用于吞吐计算。权重可放入只是基础校验，不代表服务运行时容量满足。")

    doc.add_heading("4. 理论 Decode 上限", level=1)
    add_formula(doc, "单卡理论吞吐", "单卡 BF16/FP16 FLOPS × 量化内核倍率 ÷（2 × 每 Token 激活参数）")
    doc.add_paragraph("公式中的 2 × 激活参数是 Transformer 前向 Decode 的粗略 FLOPs 近似。该值不包含访存、通信、排队、批处理和服务框架损耗，因此只作为上限。")

    doc.add_heading("5. Amdahl 非线性扩展", level=1)
    add_formula(doc, "通用加速比", "S(n) = 1 ÷ [s + (1 - s) ÷ n]", "n 为资源数量，s 为不可并行比例。s = 0 时恢复理想线性扩展；s 越大，增加资源的边际收益越低。")
    add_formula(doc, "节点内加速比合计", "Σ S（第 i 节点实际卡数, 节点内不可并行比例）")
    add_formula(doc, "跨节点通信效率", "S（单副本节点数, 跨节点不可并行比例）÷ 单副本节点数")
    add_formula(doc, "无压测单副本吞吐", "单卡理论吞吐 × 节点内加速比合计 × 跨节点通信效率 × 内核与服务效率")
    add_formula(doc, "总吞吐", "单副本吞吐 × 完整副本数 × 多副本调度效率")
    add_data_table(
        doc,
        ["模板", "服务效率", "卡内 s", "节点间 s", "定位"],
        [
            ["交互低并发", "5%", "15%", "25%", "保守敏感性"],
            ["在线均衡", "10%", "10%", "15%", "默认规划值"],
            ["吞吐优先", "20%", "5%", "10%", "优化场景参考"],
        ],
        [1900, 1600, 1400, 1600, 2860],
    )
    doc.add_paragraph("Amdahl 公式具有明确的并行计算理论依据；表中的百分比是工程规划先验，不是厂商官方倍率。")

    doc.add_heading("6. 多机和完整副本", level=1)
    add_intro_callout(doc, "副本内部署", "多个节点共同承载一个无法单机放入的模型。单副本总卡数按节点容量自动装箱，例如 19 卡、3 个 8 卡节点为 8+8+3；单副本吞吐仍按非线性扩展或压测计算。")
    add_intro_callout(doc, "同步增加硬件（默认）", "每增加一个完整副本，就增加一套对应节点和卡。总节点数等于副本数乘单副本节点数，吞吐和月租同步扩展，最容易核对。")
    add_intro_callout(doc, "使用固定节点池", "系统用“固定池节点数 ÷ 单副本节点数”自动计算每组可运行副本数，再按副本需求向上取整计算资源组数。")
    add_intro_callout(doc, "高级：实测空卡共用", "只有同模型、精度、拓扑和运行时的部署或压测已确认整张空卡可安全共用时，才手工填写每组可运行副本数；超过该数量后增加下一组硬件。")
    add_intro_callout(doc, "空卡边界", "空卡是未分配给任何副本的整张卡；已分配卡上的剩余显存不能作为空卡。即使整卡数量足够，也必须验证具体卡位、互联域、进程隔离、显存与运行时开销和调度约束。", fill="FFF4E5", color=AMBER)
    doc.add_paragraph("TP/PP/EP、互联带宽、RDMA 和推理框架不直接代入一个固定吞吐公式，而是作为配置校验和压测指纹。其精确性能仍由实测确定。")

    doc.add_heading("7. QPS 容量规划", level=1)
    add_formula(doc, "对外输出需求", "QPS × 平均输出 Token")
    add_formula(doc, "无压测等效规划负载", "QPS ×（输出 Token + 输入 Token × Prefill 等效系数）")
    add_formula(doc, "无压测单副本规划能力", "单卡理论吞吐 × 节点内加速比合计 × 跨节点通信效率 × 内核与业务规划系数")
    add_formula(doc, "服务副本数", "向上取整[规划负载 ÷（单副本能力 × 目标负载率 × 调度效率）]")
    add_formula(doc, "付费副本数", "服务副本数 + 冗余副本数")
    doc.add_paragraph("若存在严格匹配压测，系统以实测输出 TPS 直接规划，不再重复附加 Prefill 等效负载；成本分母始终是实际输出 Token。")

    doc.add_heading("8. 租金、成本与报价", level=1)
    add_formula(doc, "同步增加硬件", "总节点数 = 付费副本数 × 单副本节点数")
    add_formula(doc, "固定节点池每组可运行副本数", "向下取整[每组节点数 ÷ 单副本节点数]")
    add_formula(doc, "固定节点池资源组数", "向上取整[付费副本数 ÷ 每组可运行副本数]")
    add_formula(doc, "高级实测空卡共用", "手工回填实测值，且不得超过向下取整[每组节点数 × 单节点标准卡数 ÷ 单副本实际总卡数]")
    add_formula(doc, "整节点计费", "总月租 = 节点月租 × 每组节点数 × 资源组数")
    add_formula(doc, "按卡计费", "固定节点池：预留池卡数 × 单卡月租；同步增加硬件：单副本实际卡数 × 副本数 × 单卡月租")
    add_formula(doc, "月输出 Token", "实际输出 TPS × 运行天数 × 日均有效负载小时 × 3600")
    add_formula(doc, "硬件成本", "总月租 ÷ 月输出 Token × 1,000,000")
    add_formula(doc, "建议报价", "硬件成本 ÷（1 - 目标毛利率）")
    add_intro_callout(doc, "解释", "同步增加硬件时，资源和吞吐若同比例增长，单位 Token 成本近似不变。固定节点池按资源组阶梯计费；只有实测空卡共用确认后，才允许在每组实测容量内增加副本而月租不变。")

    doc.add_heading("9. 配置校验与压测失配", level=1)
    add_data_table(
        doc,
        ["规则", "处理方式"],
        [
            ["单机实际卡数超过硬件规格", "停用月租、成本和报价，提示最少节点数并要求切换多机"],
            ["多机单副本卡数超过拓扑容量", "停用成本和报价，要求增加单副本节点或减少卡数"],
            ["固定池节点少于单副本节点", "判定该节点池无法运行一个完整副本"],
            ["每组实测可运行副本数为空或超过整卡容量", "判定固定节点池无效，停用成本与报价"],
            ["权重超过预留后可用显存", "显示至少所需卡数/节点数，停用报价"],
            ["TP × PP 超过单副本总卡数", "判定拓扑无效"],
            ["压测工作负载或拓扑不匹配", "保留记录但不应用，回退工程估算"],
            ["压测未满足已设置 P95 SLA", "不应用该压测，回退工程估算"],
        ],
        [3600, 5760],
    )

    doc.add_heading("10. 示例：19 卡单副本与固定节点池", level=1)
    doc.add_paragraph("以 626 TFLOPS/卡、40B 激活参数、8 卡/节点、在线均衡模板为例：")
    add_formula(doc, "单卡理论上限", "626 × 10¹² ÷（2 × 40 × 10⁹）= 7,825 token/s")
    add_formula(doc, "卡分布", "19 张卡跨 3 个 8 卡节点自动分布为 8 + 8 + 3；物理容量 24 张，空余 5 张")
    add_formula(doc, "节点内加速比合计", "S(8,10%) + S(8,10%) + S(3,10%) ≈ 4.71 + 4.71 + 2.50 = 11.91×")
    add_formula(doc, "跨节点通信效率", "S(3,15%) ÷ 3 ≈ 76.92%")
    add_formula(doc, "单副本规划吞吐", "7,825 × 11.91 × 76.92% × 10% ≈ 7,170 token/s")
    add_formula(doc, "同步增加硬件", "2 个副本 × 3 节点/副本 = 6 个付费节点；月租 55,000 × 6 = 330,000 元")
    add_formula(doc, "5 节点固定池", "向下取整[5 ÷ 3] = 每组 1 个副本；规划 2 个副本需 2 组、10 个付费节点")
    add_formula(doc, "高级实测空卡共用", "理论整卡容量为向下取整[40 ÷ 19] = 2；只有部署或压测验证可共用时，才可填写每组 2 个副本并按 5 个节点计费")
    add_formula(doc, "验证后两副本总吞吐", "7,170 × 2 × 95% ≈ 13,623 token/s；5 节点月租 275,000 元，硬件成本约 7.79 元/百万 Token")
    doc.add_paragraph("该示例同时展示默认扩容、固定节点池和实测空卡共用三种结果。模型、精度、拓扑、框架、ISL/OSL 或并发变化后，应重新验证每组可运行副本数和单副本 TPS。")

    doc.add_heading("11. 依据与边界", level=1)
    sources = [
        ("Amdahl, Validity of the Single Processor Approach to Achieving Large Scale Computing Capabilities", "https://doi.org/10.1145/1465482.1465560"),
        ("Berkeley Roofline Model", "https://digicoll.lib.berkeley.edu/record/136692/files/EECS-2008-134.pdf"),
        ("NVIDIA NIM Benchmarking - Metrics", "https://docs.nvidia.com/nim/benchmarking/llm/latest/metrics.html"),
        ("NVIDIA NIM Benchmarking - Parameters", "https://docs.nvidia.com/nim/benchmarking/llm/latest/parameters.html"),
        ("vLLM - Parallelism and Scaling", "https://docs.vllm.ai/en/v0.22.0/serving/parallelism_scaling/"),
        ("vLLM - Bench Serve", "https://docs.vllm.ai/en/latest/cli/bench/serve/"),
    ]
    for idx, (title, url) in enumerate(sources, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(f"来源 {idx}：{title}\n"), 10.5, bold=True, color=DARK_BLUE)
        set_font(p.add_run(url), 9.5, color=BLUE)
    add_intro_callout(doc, "适用边界", "正式上线容量、SLA 承诺和对外定价必须使用同硬件、同拓扑、同运行时、同 ISL/OSL 与并发条件下的压测结果，并补充非硬件成本。", fill="FFF4E5", color=AMBER)

    path = OUT / "Token_API_成本测算工具_计算口径与更新报告_v3.1.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for output in (build_usage_guide(), build_logic_report()):
        print(output)
